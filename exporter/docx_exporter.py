from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

import datetime
from hashlib import md5
from base64 import b64encode

import polars as pl
from os.path import basename
import subprocess

import re


from exporter.exporter import Exporter


class DOCXExporter(Exporter):
    def __init__(self, output="docx", tmpFolder=None):
        super().__init__(tmpFolder=tmpFolder)
        self.output = output
        self.document = Document()


    def finish_page(self):
        self.document.add_page_break()

    def add_blank_page(self):
        # NO need for new pages, only add the breaks when done with "page" content
        return len(self.pages)
    
    def add_title(self, title, page_idx=None):
        if page_idx == None:
            page_idx = len(self.pages)-1

        self.document.add_heading(title, level=1)
        self.document.add_paragraph()
    
    def add_footer(self, page_idx=None):
        pass


    def add_image(self, imgInfo, page_idx=None):
        if page_idx is None:
            page_idx = len(self.pages)-1
        
        
        pg = self.document.add_paragraph().add_run()
       

        pgImg = pg.add_picture(imgInfo[0],  width=Inches(5.8))

        
        heightRel = pgImg.height / self.document.sections[0].page_height
        widthRel = pgImg.width / self.document.sections[0].page_width
       
        if heightRel >= 0.8 or widthRel >= 0.8:
            if widthRel > heightRel:
                relChange = 0.8 / widthRel
                pgImg.width = int(pgImg.width * relChange * 0.8)
                pgImg.height = int(pgImg.height * relChange * 0.8)
            else:
                relChange = 0.8 / heightRel
                pgImg.width = int(pgImg.width * relChange * 0.8)
                pgImg.height = int(pgImg.height * relChange * 0.8)
            

    def __add_table(self, tableLines, text_size=11):
        header = tableLines[0]
        nCols = len(re.findall("[|]", header)) - 1
        nRows = len(tableLines) - sum([l.startswith("|-") and l.endswith("-|") for l in tableLines])

        table = self.document.add_table(nRows, nCols)
        table.style = 'Light List'

        ri = 0
        for line in tableLines:
            if line.startswith("|-") and ":|:" in line:
                continue
            content = line.split("|")
            ci = 0
            for i in range(1, len(content)-1):
                run = table.cell(ri, ci).paragraphs[0].add_run()
                run.text =  str.strip(content[i])

                if ri == 0:
                    table.cell(ri, ci).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    table.cell(ri, ci).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run.font.size = Pt(text_size+3)
                    run.font.bold = True

                ci += 1
            ri += 1
                    

    def add_text(self, textFile, page_idx=0, text_size=11):
        with open(textFile, "r") as file:
            text = file.readlines()

        line=text[0]
        y = 1.0


        tableLines = []
        drawingTable = False
        for line in text:
            y += 0.19
            if line == "\n":
                continue

            line = line.replace("\n", "")

            if not line.startswith("|") and drawingTable == True:
                drawingTable = False
                self.__add_table(tableLines=tableLines, text_size=text_size)

            if line.startswith("#"):
                run = self.document.add_paragraph().add_run()
                run.font.bold = True
                
                if line.startswith("###"):
                    run.text =  line.split("### ")[1]
                    run.font.size = Pt(text_size + 2)

                if line.startswith("###"):
                    run.text =  line.split("## ")[1]
                    run.font.size = Pt(text_size + 4)

                if line.startswith("###"):
                    run.text =  line.split("# ")[1]
                    run.font.size = Pt(text_size + 6)
            elif line.startswith("|"):
                tableLines.append(line)

                if drawingTable == False:
                    drawingTable = True
            else:

                run = self.document.add_paragraph().add_run()
                run.text =  line
                run.font.size = Pt(text_size)
                run.font.bold = False

        if drawingTable == True:
            drawingTable = False
            self.__add_table(tableLines=tableLines, text_size=text_size)
                    

    def save(self, filename):
        self.filename = filename
        self.document.save(self.filename)
    def as_dataframe(self):
        # self.document.save("test.docx")
        filename = self.filename
        
        

        if self.output == "docx":
            outname = basename(self.filename) + ".docx"
            mimetype = "application/vnd.ms-word"
            with open(self.filename, "rb") as file:
                fileBytes = file.read()
        else:
            outname = basename(self.filename) + ".pdf"
            mimetype = "application/pdf"
            subprocess.call(["libreoffice", "--headless" ,\
                            "--convert-to", "pdf", \
                            self.filename, \
                            "--outdir", self.tmpFolder])
            with open(sself.filename.replace(".docx", ".pdf"), "rb") as file:
                fileBytes = file.read()
        
        checksum = md5(fileBytes).hexdigest()

        imgDf = pl.DataFrame({\
            ".ci":[0],\
            "filename":[outname],\
            "mimetype":[mimetype],\
            "checksum":[checksum],\
            ".content":[str(b64encode(fileBytes).decode("utf-8"))]\
            })
        
        imgDf = imgDf.with_columns(pl.col('.ci').cast(pl.Int32))
        return imgDf
